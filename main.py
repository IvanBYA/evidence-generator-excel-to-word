import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EXCEL_FILE = "Matrix_TestCases_V1.xlsx"
OUTPUT_DIR = "output"

# Columnas requeridas en el Excel, en el orden exacto esperado.
# Estas se usan para validar que el archivo seleccionado tenga el formato correcto.
REQUIRED_COLUMNS = [
    "Funcionalidad",
    "Id CP",
    "Caso de Prueba",
    "Descripcion",
    "Pre-Requisitos",
    "Datos de Prueba",
    "No. Paso",
    "Descripcion del Paso",
    "Resultado Esperado",
    "Criticidad",
]


def set_cell_bold_label(paragraph, label, value=""):
    """
    Escribe dentro de un párrafo una etiqueta en negritas seguida de un valor en texto normal.

    Se usa para mostrar campos con formato 'Etiqueta: Valor' dentro de una celda del documento Word.

    Parámetros:
        paragraph: El párrafo de Word donde se va a escribir el contenido.
        label (str): El texto de la etiqueta que aparecerá en negritas (ej. 'Funcionalidad').
        value (str): El texto del valor que aparecerá en formato normal después de la etiqueta.
                     Si se omite o está vacío, solo se escribe la etiqueta.
    """
    run_label = paragraph.add_run(label)
    run_label.bold = True
    if value:
        run_value = paragraph.add_run(value)
        run_value.bold = False


def set_font(run, size=11):
    """
    Aplica la fuente Arial y un tamaño de letra específico a un fragmento de texto (run) en Word.

    Se usa en todos los textos del documento para mantener un estilo visual uniforme.

    Parámetros:
        run: El fragmento de texto de Word al que se le aplicará el formato.
        size (int): El tamaño de la letra en puntos. Por defecto es 11pt.
    """
    run.font.size = Pt(size)
    run.font.name = "Arial"


def add_row(table, label, value, bold_value=False):
    """
    Agrega una fila simple a la tabla del documento Word con una etiqueta en negritas y su valor.

    Se usa para los campos que tienen un solo valor de texto, como 'Funcionalidad',
    'Id Caso de Prueba', 'Estatus', etc.

    Parámetros:
        table: La tabla de Word donde se va a insertar la nueva fila.
        label (str): El nombre del campo que aparecerá en negritas (ej. 'Num. Release').
        value (str): El contenido del campo. Si está vacío, se escribe '[NA]' automáticamente.
        bold_value (bool): Si es True, el valor también se mostrará en negritas. Por defecto es False.

    Retorna:
        cell: La celda de Word que fue creada, por si se necesita modificar después.
    """
    row = table.add_row()
    cell = row.cells[0]
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]

    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    set_font(label_run)

    value_run = p.add_run(str(value) if value else "[NA]")
    value_run.bold = bold_value
    set_font(value_run)

    return cell


def add_row_with_steps(table, label, steps_list):
    """
    Agrega una fila a la tabla del documento Word cuyo contenido es una lista numerada de elementos.

    Se usa para los campos 'Pasos' y 'Resultado Esperado', donde cada elemento de la lista
    se muestra en su propia línea con su número correspondiente (1. ... 2. ... 3. ...).

    Parámetros:
        table: La tabla de Word donde se va a insertar la nueva fila.
        label (str): El nombre del campo que aparecerá en negritas como encabezado de la lista
                     (ej. 'Pasos' o 'Resultado Esperado').
        steps_list (list): Lista de textos, donde cada elemento representa un paso o resultado.
                           Ejemplo: ['Abrir el navegador', 'Ir al login', 'Ingresar credenciales']

    Retorna:
        cell: La celda de Word que fue creada, por si se necesita modificar después.
    """
    row = table.add_row()
    cell = row.cells[0]
    cell.paragraphs[0].clear()

    # Label paragraph
    p_label = cell.paragraphs[0]
    label_run = p_label.add_run(f"{label}:")
    label_run.bold = True
    set_font(label_run)

    # One paragraph per step
    for i, step in enumerate(steps_list, start=1):
        p_step = cell.add_paragraph()
        step_run = p_step.add_run(f"{i}. {step}")
        set_font(step_run)

    return cell


def generate_document(case_id, funcionalidad, escenario, descripcion, pasos, resultados):
    """
    Crea y construye un documento Word completo con el formato de evidencia de pruebas.

    Esta función es la encargada de armar todo el documento: configura los márgenes de la página,
    escribe el título 'Evidencia', construye la tabla principal con todos los campos del formato
    (incluyendo los campos manuales como [NA]), y al final agrega la sección de 'Pasos' fuera
    de la tabla con el salto de línea entre cada paso.

    Parámetros:
        case_id (str): El identificador único del caso de prueba (ej. 'QA-FUNC-001').
                       Se usa también como nombre del archivo Word generado.
        funcionalidad (str): La funcionalidad que cubre el caso (ej. 'Login').
        escenario (str): El nombre del escenario de prueba (ej. 'Login exitoso...').
        descripcion (str): La descripción general del objetivo del caso de prueba.
        pasos (list): Lista de textos con cada paso a ejecutar durante la prueba.
                      Ejemplo: ['Navegar al login', 'Ingresar correo', 'Hacer clic en Entrar']
        resultados (list): Lista de textos con el resultado esperado por cada paso.
                           Debe tener la misma cantidad de elementos que 'pasos'.

    Retorna:
        doc: El objeto Document de python-docx con el documento Word completamente armado,
             listo para guardarse en disco con doc.save().
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin   = Cm(1)
    section.right_margin  = Cm(1)

    # ── Title ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title_run = title.add_run("Evidencia")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"
    title.alignment = 1  # center

    doc.add_paragraph()  # spacing

    # ── Main table ────────────────────────────────────────────────
    table = doc.add_table(rows=0, cols=1)
    table.style = "Table Grid"

    # Fixed [NA] fields
    add_row(table, "Num. Version",           "[NA]")
    add_row(table, "Funcionalidad",          funcionalidad)
    add_row(table, "Id CP",      case_id)
    add_row(table, "Caso de Prueba",    escenario)

    # Descripción (may be multiline)
    row_desc = table.add_row()
    cell_desc = row_desc.cells[0]
    cell_desc.paragraphs[0].clear()
    p_desc_label = cell_desc.paragraphs[0]
    r = p_desc_label.add_run("Descripción:")
    r.bold = True
    set_font(r)
    p_desc_val = cell_desc.add_paragraph()
    r2 = p_desc_val.add_run(str(descripcion) if descripcion else "[NA]")
    set_font(r2)

    # Pasos (numbered list inside cell)
    add_row_with_steps(table, "Pasos", pasos)

    # Resultado Esperado (numbered list inside cell)
    add_row_with_steps(table, "Resultado Esperado", resultados)

    # Remaining [NA] fields
    add_row(table, "Descripción del Defecto",
            "[NA] [Explicar mas a fondo el defecto o causa raíz]")
    add_row(table, "Estatus",                "[NA] [Successful or Failed]")
    add_row(table, "Nombre del Tester",      "[NA]")
    add_row(table, "Nombre del Developer", "[NA]")
    add_row(table, "Fecha de ejecución",     "DD-MM-YYYY")

    # ── Steps section OUTSIDE the table ───────────────────────────
    doc.add_paragraph()
    p_steps_title = doc.add_paragraph()
    r_title = p_steps_title.add_run("Pasos:")
    r_title.bold = True
    set_font(r_title, size=11)

    for i, paso in enumerate(pasos, start=1):
        p_paso = doc.add_paragraph()
        r_paso = p_paso.add_run(f"{i}. {paso}")
        set_font(r_paso)
        doc.add_paragraph()  # blank line between steps (salto de línea)

    return doc


def generate_all(excel_path, output_dir, progress_callback=None):
    """
    Función central reutilizable que lee el Excel y genera todos los documentos Word.

    A diferencia de main(), esta función está diseñada para ser llamada desde otros módulos
    (como gui.py) pasándole rutas dinámicas. Además, acepta un callback de progreso para
    que la interfaz gráfica pueda actualizar su barra en tiempo real.

    Parámetros:
        excel_path (str): Ruta completa al archivo Excel (.xlsx) seleccionado por el usuario.
        output_dir (str): Ruta de la carpeta destino donde se guardarán los documentos Word.
        progress_callback (callable | None): Función opcional que se llama después de generar
            cada documento. Recibe dos argumentos: (documentos_generados, total_documentos).
            Ejemplo: progress_callback(2, 5) indica que se generaron 2 de 5 documentos.

    Retorna:
        int: El número total de documentos Word generados exitosamente.

    Lanza:
        Exception: Si el Excel no se puede leer o si ocurre un error al guardar algún documento.
    """
    # ── Read Excel ────────────────────────────────────────────────
    df = pd.read_excel(excel_path, sheet_name="TestCases", header=1)

    # Drop completely empty rows
    df = df.dropna(how="all")

    # Forward-fill main case fields (they only appear on the first row of each case)
    fill_cols = ["Funcionalidad", "Id CP", "Caso de Prueba", "Descripcion"]
    df[fill_cols] = df[fill_cols].ffill()

    # ── Create output folder ──────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)

    # ── Group by test case and generate one Word per case ─────────
    grouped = df.groupby("Id CP", sort=False)
    total     = len(grouped)
    generated = 0

    for case_id, group in grouped:
        first_row     = group.iloc[0]
        funcionalidad = str(first_row["Funcionalidad"]).strip()
        escenario     = str(first_row["Caso de Prueba"]).strip()
        descripcion   = str(first_row["Descripcion"]).strip()

        # Collect steps and results in order
        pasos      = []
        resultados = []
        for _, row in group.iterrows():
            paso = row.get("Descripcion del Paso")
            res  = row.get("Resultado Esperado")
            if pd.notna(paso) and str(paso).strip():
                pasos.append(str(paso).strip())
            if pd.notna(res) and str(res).strip():
                resultados.append(str(res).strip())

        # Generate document
        doc = generate_document(
            case_id       = str(case_id).strip(),
            funcionalidad = funcionalidad,
            escenario     = escenario,
            descripcion   = descripcion,
            pasos         = pasos,
            resultados    = resultados,
        )

        # Safe filename: replace characters not allowed in Windows filenames
        safe_name   = str(case_id).strip().replace("/", "-").replace("\\", "-")
        output_path = os.path.join(output_dir, f"{safe_name}.docx")
        doc.save(output_path)

        generated += 1

        # Notify GUI of progress after each document is saved
        if progress_callback:
            progress_callback(generated, total)

    return generated


def main():
    """
    Función principal que orquesta todo el proceso de generación de evidencias.

    Es el punto de entrada del programa cuando se ejecuta desde la línea de comandos.
    Valida que el Excel exista y delega toda la lógica de generación a generate_all().

    No recibe parámetros ni retorna valores. Si el Excel no existe o no se puede leer,
    el programa muestra un mensaje de error y se detiene.
    """
    # ── Validate Excel exists ─────────────────────────────────────
    if not os.path.exists(EXCEL_FILE):
        print(f"[ERROR] No se encontró el archivo '{EXCEL_FILE}' en el directorio actual.")
        print(f"        Asegúrate de ejecutar el script en la misma carpeta que el Excel.")
        sys.exit(1)

    # ── Generate all documents via shared function ────────────────
    try:
        total = generate_all(
            excel_path        = EXCEL_FILE,
            output_dir        = OUTPUT_DIR,
            progress_callback = lambda done, t: print(f"[OK] Generado {done}/{t}"),
        )
    except Exception as e:
        print(f"[ERROR] No se pudo completar la generación: {e}")
        sys.exit(1)

    print(f"\n✔  {total} documento(s) generado(s) en la carpeta '{OUTPUT_DIR}'")


if __name__ == "__main__":
    main()
